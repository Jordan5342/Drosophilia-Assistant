import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x2E, 0x5F, 0xA3)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x16, 0x65, 0x34)
AMBER = RGBColor(0x92, 0x40, 0x0E)


def _run(para, text, size_pt, color, bold=False):
    r = para.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size_pt)
    r.font.color.rgb = color
    r.font.bold = bold
    return r


def _spacing(para, before_pt, after_pt):
    para.paragraph_format.space_before = Pt(before_pt)
    para.paragraph_format.space_after = Pt(after_pt)


def _bottom_border(para, color_hex, size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _new_doc():
    doc = Document()
    for section in doc.sections:
        section.page_width   = Inches(8.5)
        section.page_height  = Inches(11)
        section.left_margin  = section.right_margin  = Inches(1)
        section.top_margin   = section.bottom_margin = Inches(1)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    # Remove default paragraph spacing from Normal
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    return doc


def _h1(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 18, 6)
    _bottom_border(p, '2E5FA3', 6)
    _run(p, text, 14, BLUE, bold=True)
    return p


def _h2(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 12, 4)
    _run(p, text, 12, DARK, bold=True)
    return p


def _body(doc, text, color=None):
    color = color or GRAY
    p = doc.add_paragraph()
    _spacing(p, 3, 6)
    parts = re.split(r'\*\*([^*]+)\*\*', text or '')
    for i, part in enumerate(parts):
        _run(p, part, 11, color, bold=(i % 2 == 1))
    return p


def _kv(doc, label, value, value_color=None):
    value_color = value_color or GRAY
    p = doc.add_paragraph()
    _spacing(p, 2, 2)
    _run(p, label + ': ', 11, DARK, bold=True)
    _run(p, value or '', 11, value_color)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    _spacing(p, 2, 2)
    # Clear default run added by add_paragraph and add our own styled run
    for run in p.runs:
        run.text = ''
    _run(p, text or '', 11, GRAY)
    return p


def _spacer(doc):
    p = doc.add_paragraph()
    _spacing(p, 4, 4)
    return p


def _centered(doc, text, size_pt, color, bold=False, before_pt=0, after_pt=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before_pt, after_pt)
    _run(p, text, size_pt, color, bold=bold)
    return p


def _hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_r.append(t)
    hl.append(new_r)
    para._p.append(hl)


def build_proposal_docx(proposal) -> bytes:
    doc = _new_doc()

    _centered(doc, proposal.get('title') or 'Research Proposal', 20, DARK, bold=True, before_pt=24, after_pt=12)
    _centered(doc, 'Drosophila Research Proposal', 12, BLUE, after_pt=6)
    _centered(doc, datetime.now().strftime('%B %d, %Y'), 11, GRAY, after_pt=24)
    _spacer(doc)

    _h1(doc, 'Background & Significance')
    _body(doc, proposal.get('background') or '')
    _spacer(doc)

    _h1(doc, 'Central Hypothesis')
    p = doc.add_paragraph()
    _spacing(p, 4, 4)
    _run(p, 'H\u2081: ', 11, BLUE, bold=True)
    _run(p, proposal.get('central_hypothesis') or '', 11, DARK)

    if proposal.get('null_hypothesis'):
        p = doc.add_paragraph()
        _spacing(p, 4, 4)
        _run(p, 'H\u2080: ', 11, GRAY, bold=True)
        _run(p, proposal['null_hypothesis'], 11, GRAY)

    if proposal.get('rationale'):
        _spacer(doc)
        _h2(doc, 'Rationale')
        _body(doc, proposal['rationale'])
    _spacer(doc)

    aims = proposal.get('specific_aims') or []
    if aims:
        _h1(doc, 'Specific Aims')
        for aim in aims:
            _h2(doc, f"Aim {aim.get('aim_number')}: {aim.get('title')}")
            if aim.get('objective'):        _kv(doc, 'Objective', aim['objective'])
            if aim.get('approach'):         _kv(doc, 'Approach', aim['approach'])
            if aim.get('expected_outcomes'): _kv(doc, 'Expected Outcomes', aim['expected_outcomes'])
            if aim.get('potential_pitfalls'): _kv(doc, 'Potential Pitfalls', aim['potential_pitfalls'])
            if aim.get('mitigation'):       _kv(doc, 'Mitigation', aim['mitigation'])
            _spacer(doc)

    if proposal.get('experimental_approach'):
        _h1(doc, 'Experimental Approach')
        _body(doc, proposal['experimental_approach'])
        _spacer(doc)

    if proposal.get('controls'):
        _h1(doc, 'Controls')
        _body(doc, proposal['controls'])
        _spacer(doc)

    gaps = proposal.get('literature_gaps') or []
    if gaps:
        _h1(doc, '\u26a0\ufe0f Literature Gaps & Research Opportunities')
        for gap in gaps:
            _bullet(doc, gap)
        _spacer(doc)

    timeline = proposal.get('timeline') or []
    if timeline:
        _h1(doc, 'Timeline')
        for phase in timeline:
            p = doc.add_paragraph()
            _spacing(p, 4, 2)
            _run(p, (phase.get('phase') or '') + ': ', 11, BLUE, bold=True)
            _run(p, phase.get('milestones') or '', 11, GRAY)
        _spacer(doc)

    refs = proposal.get('references') or []
    if refs:
        _h1(doc, 'References')
        for ref in refs:
            p = doc.add_paragraph()
            _spacing(p, 2, 2)
            _run(p, f"[{ref.get('number')}] ", 10, BLUE, bold=True)
            _run(p, ref.get('citation') or '', 10, GRAY)
            if ref.get('url'):
                _run(p, ' ', 10, GRAY)
                _hyperlink(p, '[Link]', ref['url'])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_experiment_design_docx(design) -> bytes:
    doc = _new_doc()

    _centered(doc, f"Aim {design.get('aim_number')}: {design.get('aim_title')}", 18, DARK, bold=True, before_pt=24, after_pt=6)
    _centered(doc, 'Experiment Design Protocol', 12, BLUE, after_pt=6)
    if design.get('estimated_duration'):
        _centered(doc, '\u23f1 ' + design['estimated_duration'], 11, GRAY, after_pt=24)

    cs = design.get('cross_scheme') or {}
    if cs.get('overview'):
        _h1(doc, '\U0001f9ec Cross Scheme')
        _body(doc, cs['overview'])
        if cs.get('parental_genotypes'):
            _h2(doc, 'Parental Lines')
            for pg in cs['parental_genotypes']:
                p = doc.add_paragraph()
                _spacing(p, 3, 1)
                _run(p, (pg.get('line') or '') + ': ', 11, DARK, bold=True)
                _run(p, (pg.get('genotype') or '') + ' \u2014 ' + (pg.get('source') or ''), 11, GRAY)
                if pg.get('notes'):
                    _body(doc, '  ' + pg['notes'], BLUE)
        if cs.get('generations'):
            _h2(doc, 'Crossing Generations')
            for g in cs['generations']:
                p = doc.add_paragraph()
                _spacing(p, 4, 1)
                _run(p, (g.get('generation') or '') + ': ' + (g.get('cross') or ''), 11, DARK, bold=True)
                if g.get('instructions'): _body(doc, g['instructions'])
                if g.get('timing'):       _body(doc, '\u23f1 ' + g['timing'], BLUE)
        if cs.get('expected_progeny'):
            _h2(doc, 'Expected Progeny')
            _body(doc, cs['expected_progeny'])
        if cs.get('balancer_notes'):
            _h2(doc, 'Balancer Notes')
            _body(doc, cs['balancer_notes'])
        _spacer(doc)

    vs = design.get('vial_setup') or {}
    vs_entries = [
        (label, vs.get(key)) for label, key in [
            ('Vial Type', 'vial_type'), ('Fly Density', 'fly_density'), ('Food', 'food'),
            ('Temperature', 'temperature'), ('Light Cycle', 'light_cycle'), ('Humidity', 'humidity'),
            ('Flipping Schedule', 'flipping_schedule'), ('Special Conditions', 'special_conditions'),
        ] if vs.get(key)
    ]
    if vs_entries:
        _h1(doc, '\U0001fad9 Vial Setup & Husbandry')
        for label, val in vs_entries:
            _kv(doc, label, val)
        _spacer(doc)

    tps = design.get('timepoints') or []
    if tps:
        _h1(doc, '\U0001f4c5 Timepoints & Schedule')
        for tp in tps:
            p = doc.add_paragraph()
            _spacing(p, 4, 1)
            _run(p, f"Day {tp.get('day')}: {tp.get('action')}", 11, DARK, bold=True)
            if tp.get('what_to_collect'): _body(doc, '\U0001f4cb ' + tp['what_to_collect'], BLUE)
            if tp.get('notes'):           _body(doc, tp['notes'])
        _spacer(doc)

    dcs = design.get('data_collection_sheet') or {}
    if dcs.get('columns'):
        _h1(doc, '\U0001f4ca Data Collection Sheet')
        if dcs.get('description'):      _body(doc, dcs['description'])
        _body(doc, 'Columns: ' + ' | '.join(dcs['columns']))
        if dcs.get('scoring_criteria'): _body(doc, 'Scoring: ' + dcs['scoring_criteria'], BLUE)
        _spacer(doc)

    sa = design.get('statistical_analysis') or {}
    if sa.get('primary_test'):
        _h1(doc, '\U0001f4c8 Statistical Analysis')
        if sa.get('primary_test'):              _kv(doc, 'Primary Test', sa['primary_test'])
        if sa.get('software'):                  _kv(doc, 'Software', sa['software'])
        if sa.get('sample_size_justification'): _kv(doc, 'Sample Size', sa['sample_size_justification'])
        if sa.get('censoring_criteria'):        _kv(doc, 'Censoring Criteria', sa['censoring_criteria'])
        if sa.get('multiple_comparisons'):      _kv(doc, 'Multiple Comparisons', sa['multiple_comparisons'])
        _spacer(doc)

    gng = design.get('go_no_go_criteria') or []
    if gng:
        _h1(doc, '\U0001f6a6 Go/No-Go Criteria')
        for g in gng:
            p = doc.add_paragraph()
            _spacing(p, 4, 1)
            _run(p, f"{g.get('checkpoint')} (Check: {g.get('timing')})", 11, DARK, bold=True)
            _body(doc, '\u2705 Go if: ' + (g.get('threshold') or ''), GREEN)
            _body(doc, '\u26a0\ufe0f If fail: ' + (g.get('if_fail') or ''), AMBER)
        _spacer(doc)

    ts = design.get('troubleshooting') or []
    if ts:
        _h1(doc, '\U0001f527 Troubleshooting')
        for t in ts:
            p = doc.add_paragraph()
            _spacing(p, 4, 1)
            _run(p, 'Problem: ' + (t.get('problem') or ''), 11, DARK, bold=True)
            _body(doc, 'Cause: ' + (t.get('likely_cause') or ''))
            _body(doc, 'Solution: ' + (t.get('solution') or ''), BLUE)
        _spacer(doc)

    mats = design.get('materials_needed') or []
    if mats:
        _h1(doc, '\U0001f9ea Materials Needed')
        for m in mats:
            _bullet(doc, m)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
